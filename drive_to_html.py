#!/usr/bin/env python3
"""
drive_to_html.py — Lüks İmmersive HTML Sunum Üretici (v2)
────────────────────────────────────────────────────────────
Google Drive klasöründeki TÜM dosyaları okuyup
premium, sinematik bir HTML sunum sayfası üretir.

Desteklenen dosya türleri:
  • Görseller   : jpg, jpeg, png, gif, webp, svg
  • PDF         : canvas tabanlı viewer + sayfa thumbnails
  • Word        : .docx paragraf / başlıklar
  • Excel/CSV   : tablo verisi
  • JSON        : yapılandırılmış veri
  • Diğerleri   : metadata kartı

Kurulum:
  pip install google-api-python-client google-auth-httplib2 \
              google-auth-oauthlib pandas openpyxl python-docx \
              PyPDF2 Pillow

OAuth credentials.json dosyasını Google Cloud Console'dan indirip
scriptle aynı dizine koyun.
────────────────────────────────────────────────────────────
"""

import os, io, json, base64, textwrap, html, re, hashlib, pickle, shutil
from pathlib import Path
from datetime import datetime

# ── Google API ──────────────────────────────────────────────
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload

# ── Dosya işleme ────────────────────────────────────────────
import pandas as pd
from docx import Document as DocxDocument
import PyPDF2
from PIL import Image

# ════════════════════════════════════════════════════════════
#  AYARLAR
# ════════════════════════════════════════════════════════════
FOLDER_ID        = os.environ.get("DRIVE_FOLDER_ID", "1wl6IORLksewhrWqpCOfjFNgjlC_rAhZT")
SCOPES           = ["https://www.googleapis.com/auth/drive.readonly"]
OUTPUT_FILE      = "sunum.html"
PROJE_ADI        = "Proje Sunumu"
PROJE_ALT_BASLIK = "Google Drive Arşivi"
CACHE_DIR        = ".drive_cache"
ASSETS_DIR       = "sunum_assets"


# ════════════════════════════════════════════════════════════
#  1. GOOGLE DRIVE KİMLİK DOĞRULAMA
# ════════════════════════════════════════════════════════════
def get_service():
    """
    Render.com için Service Account ile kimlik doğrulama.
    GOOGLE_SERVICE_ACCOUNT_JSON env var'ı gereklidir.
    """
    sa_json_str = os.environ.get("GOOGLE_SERVICE_ACCOUNT_JSON")
    if not sa_json_str:
        raise EnvironmentError(
            "GOOGLE_SERVICE_ACCOUNT_JSON ortam değişkeni bulunamadı!\n"
            "Render Dashboard → Environment → GOOGLE_SERVICE_ACCOUNT_JSON\n"
            "Değer olarak Service Account JSON dosyasının tüm içeriğini yapıştırın."
        )
    sa_info = json.loads(sa_json_str)
    creds = service_account.Credentials.from_service_account_info(
        sa_info, scopes=SCOPES
    )
    return build("drive", "v3", credentials=creds)


# ════════════════════════════════════════════════════════════
#  2. DOSYA LİSTESİ ÇEK (özyinelemeli)
# ════════════════════════════════════════════════════════════
def list_files(service, folder_id, _depth=0, _path=""):
    results, page_token = [], None
    while True:
        resp = service.files().list(
            q=f"'{folder_id}' in parents and trashed=false",
            spaces="drive",
            fields="nextPageToken, files(id,name,mimeType,size,modifiedTime)",
            pageToken=page_token
        ).execute()
        items = resp.get("files", [])
        for item in items:
            if item.get("mimeType") == "application/vnd.google-apps.folder":
                sub_path = (_path + " / " if _path else "") + item["name"]
                print(f"  {'  ' * _depth}📁 {item['name']} (alt klasör, içine giriliyor…)")
                sub_files = list_files(service, item["id"], _depth + 1, sub_path)
                results.extend(sub_files)
            else:
                item["folder_path"] = _path
                results.append(item)
        page_token = resp.get("nextPageToken")
        if not page_token:
            break
    return results


# ════════════════════════════════════════════════════════════
#  3. DOSYA İNDİR
# ════════════════════════════════════════════════════════════
GAPPS_EXPORT = {
    "application/vnd.google-apps.presentation": ("application/pdf", "pdf"),
    "application/vnd.google-apps.document":     ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"),
    "application/vnd.google-apps.spreadsheet":  ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "xlsx"),
    "application/vnd.google-apps.drawing":      ("image/png", "png"),
}

def download_bytes(service, file_id, mime_type="") -> tuple[bytes, str]:
    def _dl(req) -> bytes:
        buf = io.BytesIO()
        dl = MediaIoBaseDownload(buf, req)
        done = False
        while not done:
            _, done = dl.next_chunk()
        return buf.getvalue()

    if mime_type in GAPPS_EXPORT:
        export_mime, _ = GAPPS_EXPORT[mime_type]
        return _dl(service.files().export_media(fileId=file_id, mimeType=export_mime)), export_mime

    try:
        return _dl(service.files().get_media(fileId=file_id)), mime_type
    except Exception as e:
        if "fileNotDownloadable" not in str(e) and "403" not in str(e):
            raise

    for export_mime in [
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.presentationml.presentation",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "image/png",
    ]:
        try:
            data = _dl(service.files().export_media(fileId=file_id, mimeType=export_mime))
            return data, export_mime
        except Exception:
            continue

    raise RuntimeError(f"Dosya indirilemedi: file_id={file_id}, mime={mime_type}")


# ════════════════════════════════════════════════════════════
#  3b. ÖNBELLEK
# ════════════════════════════════════════════════════════════
def _cache_key(file_id: str, modified_time: str) -> str:
    raw = f"{file_id}_{modified_time}".encode()
    return hashlib.sha1(raw).hexdigest()

def cache_get(file_id: str, modified_time: str):
    key  = _cache_key(file_id, modified_time)
    path = Path(CACHE_DIR) / key
    if path.exists():
        try:
            with open(path, "rb") as f:
                return pickle.load(f)
        except Exception:
            path.unlink(missing_ok=True)
    return None

def cache_set(file_id: str, modified_time: str, data: bytes, mime: str):
    Path(CACHE_DIR).mkdir(exist_ok=True)
    key  = _cache_key(file_id, modified_time)
    path = Path(CACHE_DIR) / key
    with open(path, "wb") as f:
        pickle.dump((data, mime), f)

def cache_stats() -> dict:
    d = Path(CACHE_DIR)
    if not d.exists():
        return {"files": 0, "size_mb": 0.0}
    files = list(d.iterdir())
    total = sum(p.stat().st_size for p in files if p.is_file())
    return {"files": len(files), "size_mb": total / (1024*1024)}

def download_cached(service, file_id: str, mime_type: str, modified_time: str) -> tuple[bytes, str]:
    hit = cache_get(file_id, modified_time)
    if hit:
        return hit
    data, real_mime = download_bytes(service, file_id, mime_type)
    cache_set(file_id, modified_time, data, real_mime)
    return data, real_mime


# ════════════════════════════════════════════════════════════
#  4. DOSYA TİPİNE GÖRE İŞLE
# ════════════════════════════════════════════════════════════
def ext(name: str) -> str:
    return Path(name).suffix.lower().lstrip(".")

def size_fmt(b) -> str:
    if b is None: return "—"
    b = int(b)
    for unit in ("B","KB","MB","GB"):
        if b < 1024: return f"{b:.1f} {unit}"
        b /= 1024
    return f"{b:.1f} TB"

# ── PDF (global modal tabanlı) ──────────────────────────────
def process_pdf(data: bytes, name: str, label: str = "PDF") -> str:
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(data))
        pages  = len(reader.pages)
    except Exception:
        pages = 0
    uid      = hashlib.md5(data[:128]).hexdigest()[:12]
    pdf_path = save_pdf(data, uid)
    lbl      = "🎞 Sunum" if label == "Slides" else "📄 PDF"
    pg_label = f"{pages} Sayfa" if pages else "PDF"
    dl_name  = html.escape(name)

    return f"""<div class="pdf-card" onclick="openPdfModal('{pdf_path}','{dl_name}',{pages})">
  <div class="pdf-card-preview">
    <canvas class="pdf-thumb-canvas" data-pdf="{pdf_path}" data-uid="pthumb-{uid}"></canvas>
    <div class="pdf-card-overlay">
      <div class="pdf-open-icon">
        <svg width="28" height="28" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5"><path d="M14 2H6a2 2 0 00-2 2v16a2 2 0 002 2h12a2 2 0 002-2V8z"/><path d="M14 2v6h6"/><line x1="16" y1="13" x2="8" y2="13"/><line x1="16" y1="17" x2="8" y2="17"/><line x1="10" y1="9" x2="8" y2="9"/></svg>
      </div>
      <span class="pdf-open-label">Görüntüle</span>
    </div>
  </div>
  <div class="pdf-card-info">
    <div class="pdf-card-tag">{lbl}</div>
    <div class="pdf-card-name">{dl_name}</div>
    <div class="pdf-card-meta">{pg_label}</div>
  </div>
</div>"""

# ── Word ─────────────────────────────────────────────────────
def process_docx(data: bytes, name: str) -> str:
    try:
        doc  = DocxDocument(io.BytesIO(data))
        rows = []
        for para in doc.paragraphs[:40]:
            t = para.text.strip()
            if not t: continue
            style = para.style.name.lower()
            if "heading 1" in style:
                rows.append(f"<h4 class='doc-h1'>{html.escape(t)}</h4>")
            elif "heading 2" in style:
                rows.append(f"<h5 class='doc-h2'>{html.escape(t)}</h5>")
            elif "heading 3" in style:
                rows.append(f"<h5 class='doc-h3'>{html.escape(t)}</h5>")
            else:
                rows.append(f"<p class='doc-p'>{html.escape(t[:300])}{'…' if len(t)>300 else ''}</p>")
        body = "\n".join(rows) or "<p class='muted'>İçerik bulunamadı.</p>"
    except Exception as e:
        body = f"<p class='muted'>Hata: {e}</p>"
    return (f'<div class="card card-word"><div class="card-label">📝 Word</div>'
            f'<h3 class="card-title">{html.escape(name)}</h3>'
            f'<div class="doc-body">{body}</div></div>')

# ── Excel / CSV ──────────────────────────────────────────────
def process_table(data: bytes, name: str) -> str:
    try:
        e = ext(name)
        if e == "csv":
            df = pd.read_csv(io.BytesIO(data), nrows=50)
        else:
            df = pd.read_excel(io.BytesIO(data), nrows=50)
        rows_total = len(df)
        df = df.head(20)
        thead = "<tr>" + "".join(f"<th>{html.escape(str(c))}</th>" for c in df.columns) + "</tr>"
        tbody = ""
        for _, row in df.iterrows():
            tbody += "<tr>" + "".join(f"<td>{html.escape(str(v))}</td>" for v in row) + "</tr>"
        note = f"<p class='table-note'>İlk 20 satır · Toplam: {rows_total}</p>" if rows_total>20 else ""
    except Exception as e:
        thead, tbody, note = "", "", f"<p class='muted'>Hata: {e}</p>"
    return (f'<div class="card card-table"><div class="card-label">📊 {"Excel" if ext(name)!="csv" else "CSV"}</div>'
            f'<h3 class="card-title">{html.escape(name)}</h3>'
            f'<div class="table-wrap"><table><thead>{thead}</thead><tbody>{tbody}</tbody></table></div>'
            f'{note}</div>')

# ── JSON ─────────────────────────────────────────────────────
def process_json(data: bytes, name: str) -> str:
    try:
        obj = json.loads(data.decode("utf-8", errors="replace"))
        pretty = json.dumps(obj, ensure_ascii=False, indent=2)
        snippet = pretty[:1200] + ("\n…" if len(pretty) > 1200 else "")
    except Exception as e:
        snippet = f"Hata: {e}"
    return (f'<div class="card card-json"><div class="card-label">⚙ JSON</div>'
            f'<h3 class="card-title">{html.escape(name)}</h3>'
            f'<pre class="json-pre">{html.escape(snippet)}</pre></div>')

# ── Diğer ────────────────────────────────────────────────────
def process_other(meta: dict) -> str:
    folder_path = meta.get("folder_path", "")
    folder_tag = (f'<p class="meta">📁 {html.escape(folder_path)}</p>' if folder_path else "")
    return (f'<div class="card card-other"><div class="card-label">📎 Dosya</div>'
            f'<h3 class="card-title">{html.escape(meta["name"])}</h3>'
            f'{folder_tag}'
            f'<p class="meta">Boyut: {size_fmt(meta.get("size"))}</p>'
            f'<p class="muted">Önizleme desteklenmiyor.</p></div>')


# ════════════════════════════════════════════════════════════
#  5. YARDIMCI
# ════════════════════════════════════════════════════════════
def slugify(s: str) -> str:
    s = s.upper()
    tr = str.maketrans("ÇĞİÖŞÜçğıöşü", "CGIOSUcgiosu")
    s = s.translate(tr)
    return re.sub(r"[^A-Z0-9]+", "_", s).strip("_")

def file_type_key(name: str, mime: str) -> str:
    e = ext(name)
    if "google-apps.presentation" in mime: return "pdf"
    if "google-apps.document"     in mime: return "word"
    if "google-apps.spreadsheet"  in mime: return "table"
    if "google-apps.drawing"      in mime: return "image"
    if "pdf"            in mime: return "pdf"
    if "presentationml" in mime: return "pdf"
    if "wordprocessing" in mime: return "word"
    if "spreadsheetml"  in mime: return "table"
    if e in ("jpg","jpeg","png","gif","webp","svg") or mime.startswith("image/"): return "image"
    if e in ("xlsx","xls","csv"): return "table"
    if e == "docx":  return "word"
    if e == "json":  return "json"
    if e == "pdf":   return "pdf"
    return "other"


# ════════════════════════════════════════════════════════════
#  6. ASSET DOSYALARI
# ════════════════════════════════════════════════════════════
def _prepare_image(data: bytes, size: tuple) -> tuple:
    img = Image.open(io.BytesIO(data))
    img.thumbnail(size, Image.LANCZOS)
    if img.mode in ("RGBA", "LA", "P"):
        bg = Image.new("RGB", img.size, (4, 4, 12))
        paste_img = img.convert("RGBA") if img.mode != "RGBA" else img
        bg.paste(paste_img, mask=paste_img.split()[3])
        img = bg
    elif img.mode != "RGB":
        img = img.convert("RGB")
    return img

def make_thumb(data: bytes, mime: str, uid: str, size=(600, 400)) -> str:
    try:
        Path(ASSETS_DIR).mkdir(exist_ok=True)
        p = Path(ASSETS_DIR) / f"{uid}_t.jpg"
        if p.exists():
            return f"{ASSETS_DIR}/{uid}_t.jpg"
        img = _prepare_image(data, size)
        img.save(str(p), format="JPEG", quality=82, optimize=True)
        return f"{ASSETS_DIR}/{uid}_t.jpg"
    except Exception as e:
        print(f"    ⚠  Thumbnail hatası ({uid}): {e}")
        return ""

def make_large(data: bytes, mime: str, uid: str, size=(1600, 1200)) -> str:
    try:
        Path(ASSETS_DIR).mkdir(exist_ok=True)
        p = Path(ASSETS_DIR) / f"{uid}_l.jpg"
        if p.exists():
            return f"{ASSETS_DIR}/{uid}_l.jpg"
        img = _prepare_image(data, size)
        img.save(str(p), format="JPEG", quality=90, optimize=True)
        return f"{ASSETS_DIR}/{uid}_l.jpg"
    except Exception as e:
        print(f"    ⚠  Large görsel hatası ({uid}): {e}")
        return ""

def save_video(data: bytes, uid: str) -> str:
    try:
        Path(ASSETS_DIR).mkdir(exist_ok=True)
        p = Path(ASSETS_DIR) / f"{uid}.mp4"
        p.write_bytes(data)
        return f"{ASSETS_DIR}/{uid}.mp4"
    except Exception:
        return ""

def save_pdf(data: bytes, uid: str) -> str:
    try:
        Path(ASSETS_DIR).mkdir(exist_ok=True)
        p = Path(ASSETS_DIR) / f"{uid}.pdf"
        p.write_bytes(data)
        return f"{ASSETS_DIR}/{uid}.pdf"
    except Exception:
        return ""


# ════════════════════════════════════════════════════════════
#  8. MANIFEST
# ════════════════════════════════════════════════════════════
MANIFEST_FILE = ".drive_manifest"

def compute_manifest(files: list) -> str:
    parts = sorted(f"{f['id']}:{f.get('modifiedTime','')}" for f in files)
    return hashlib.sha256("\n".join(parts).encode()).hexdigest()

def load_manifest() -> str:
    p = Path(MANIFEST_FILE)
    return p.read_text(encoding="utf-8").strip() if p.exists() else ""

def save_manifest(h: str):
    Path(MANIFEST_FILE).write_text(h, encoding="utf-8")


# ════════════════════════════════════════════════════════════
#  9. HTML ŞABLONLARI — YENİDEN TASARIM
# ════════════════════════════════════════════════════════════

def _html_head(title: str) -> str:
    return f"""<!DOCTYPE html>
<html lang="tr">
<head>
<meta charset="UTF-8"/>
<meta name="viewport" content="width=device-width,initial-scale=1"/>
<title>{html.escape(title)}</title>
<link rel="preconnect" href="https://fonts.googleapis.com"/>
<link rel="stylesheet" href="https://fonts.googleapis.com/css2?family=Cormorant+Garamond:ital,wght@0,300;0,400;0,600;1,300;1,400&family=Plus+Jakarta+Sans:wght@300;400;500;600&family=DM+Mono:wght@300;400&display=swap"/>
<script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.min.js"></script>
<style>
"""


def _html_css() -> str:
    return """
/* ─── TOKENS ──────────────────────────────────────────── */
:root {
  --ink:      #020209;
  --ink2:     #08080F;
  --ink3:     #0D0D1A;
  --ink4:     #131320;
  --ink5:     #1A1A2E;
  --gold:     #C8A55A;
  --gold2:    #E0C07A;
  --gold3:    #F0D9A0;
  --gold-dim: rgba(200,165,90,.08);
  --gold-glow:rgba(200,165,90,.18);
  --plat:     #EDE8E0;
  --plat2:    #A8A4B0;
  --plat3:    #6A677A;
  --border:   rgba(200,165,90,.12);
  --border2:  rgba(200,165,90,.28);
  --border3:  rgba(200,165,90,.45);
  --r:        12px;
  --r2:       20px;
  --serif:    "Cormorant Garamond", Georgia, serif;
  --sans:     "Plus Jakarta Sans", system-ui, sans-serif;
  --mono:     "DM Mono", monospace;
  --ease:     cubic-bezier(.4,0,.2,1);
  --ease-out: cubic-bezier(0,.8,.2,1);
  --t:        .3s;
  --t2:       .55s;
}

/* ─── RESET ──────────────────────────────────────────── */
*,*::before,*::after { box-sizing:border-box; margin:0; padding:0 }
html { scroll-behavior:smooth; -webkit-font-smoothing:antialiased }
body {
  background:var(--ink);
  color:var(--plat);
  font-family:var(--sans);
  font-size:15px;
  line-height:1.7;
  min-height:100vh;
  overflow-x:hidden;
}

/* ─── LOADER ─────────────────────────────────────────── */
#loader {
  position:fixed; inset:0; z-index:9999;
  background:var(--ink);
  display:flex; align-items:center; justify-content:center;
  flex-direction:column; gap:32px;
  transition:opacity .8s var(--ease), visibility .8s;
}
#loader.hidden { opacity:0; visibility:hidden; pointer-events:none }
.loader-logo {
  font-family:var(--serif);
  font-size:clamp(2rem,5vw,3.5rem);
  font-weight:300;
  letter-spacing:.12em;
  color:var(--gold);
  opacity:0;
  animation:loaderFadeIn 1s var(--ease) .3s forwards;
}
.loader-line {
  width:80px; height:1px;
  background:linear-gradient(90deg, transparent, var(--gold), transparent);
  animation:loaderExpand 1.2s var(--ease) .5s forwards;
  transform:scaleX(0);
  transform-origin:center;
}
.loader-sub {
  font-family:var(--mono);
  font-size:.65rem;
  letter-spacing:.3em;
  text-transform:uppercase;
  color:var(--plat3);
  opacity:0;
  animation:loaderFadeIn .8s var(--ease) .9s forwards;
}

/* ─── CURSOR ─────────────────────────────────────────── */
#cursor {
  position:fixed; z-index:9998; pointer-events:none;
  width:8px; height:8px;
  background:var(--gold);
  border-radius:50%;
  transform:translate(-50%,-50%);
  transition:width .2s, height .2s, background .2s;
  mix-blend-mode:screen;
}
#cursor-ring {
  position:fixed; z-index:9997; pointer-events:none;
  width:36px; height:36px;
  border:1px solid rgba(200,165,90,.5);
  border-radius:50%;
  transform:translate(-50%,-50%);
  transition:transform .12s var(--ease), width .3s, height .3s, border-color .3s;
}
body:has(.proj-card:hover) #cursor,
body:has(.gal-item:hover) #cursor { width:16px; height:16px }
body:has(.proj-card:hover) #cursor-ring,
body:has(.gal-item:hover) #cursor-ring { width:60px; height:60px; border-color:var(--gold) }
@media (hover:none) { #cursor,#cursor-ring { display:none } }

/* ─── SCROLL PROGRESS ────────────────────────────────── */
#progress-bar {
  position:fixed; top:0; left:0; z-index:9996;
  height:2px;
  background:linear-gradient(90deg, var(--gold), var(--gold2));
  width:0%;
  transition:width .1s linear;
  box-shadow:0 0 12px var(--gold-glow);
}

/* ─── AMBIENT ORBS ───────────────────────────────────── */
.ambient-orb {
  position:fixed; pointer-events:none; z-index:0;
  border-radius:50%;
  filter:blur(80px);
  opacity:.06;
  animation:orbFloat 20s ease-in-out infinite;
}
.orb1 { width:600px; height:600px; background:var(--gold); top:-200px; right:-200px; animation-delay:0s }
.orb2 { width:400px; height:400px; background:#6040C0; bottom:-100px; left:-100px; animation-delay:-7s }
.orb3 { width:300px; height:300px; background:var(--gold2); bottom:40%; right:10%; animation-delay:-13s }

/* ─── HEADER ─────────────────────────────────────────── */
.site-header {
  position:relative; z-index:10;
  padding:80px 72px 64px;
  background:linear-gradient(170deg, #0A0A18 0%, #06060E 60%, var(--ink) 100%);
  border-bottom:1px solid var(--border);
  overflow:hidden;
}
.site-header::before {
  content:"";
  position:absolute; inset:0;
  background:radial-gradient(ellipse 80% 60% at 80% 50%, rgba(200,165,90,.06), transparent);
  pointer-events:none;
}
.header-inner { max-width:1400px; margin:0 auto; position:relative; z-index:1 }
.header-eyebrow {
  font-family:var(--mono);
  font-size:.65rem;
  letter-spacing:.35em;
  text-transform:uppercase;
  color:var(--gold);
  margin-bottom:24px;
  opacity:0;
  animation:slideUp .8s var(--ease) 1.4s forwards;
}
.site-header h1 {
  font-family:var(--serif);
  font-size:clamp(3rem,7vw,6rem);
  font-weight:300;
  letter-spacing:-.02em;
  line-height:1.05;
  color:var(--plat);
  opacity:0;
  animation:slideUp .9s var(--ease) 1.6s forwards;
}
.site-header h1 em { font-style:italic; color:var(--gold2) }
.header-sub {
  margin-top:20px;
  color:var(--plat2);
  font-size:1rem;
  font-weight:300;
  letter-spacing:.04em;
  opacity:0;
  animation:slideUp .8s var(--ease) 1.8s forwards;
}
.header-line {
  width:120px; height:1px;
  background:linear-gradient(90deg, var(--gold), transparent);
  margin:32px 0;
  opacity:0;
  animation:slideUp .8s var(--ease) 2s forwards;
}
.header-meta {
  display:flex; gap:20px; flex-wrap:wrap; align-items:center;
  opacity:0;
  animation:slideUp .8s var(--ease) 2.1s forwards;
}
.hm-pill {
  font-family:var(--mono); font-size:.65rem; letter-spacing:.12em;
  color:var(--plat3); padding:7px 20px;
  border:1px solid var(--border); border-radius:999px;
  background:rgba(255,255,255,.02);
  transition:border-color var(--t), color var(--t);
}
.hm-pill:hover { border-color:var(--border2); color:var(--plat) }
.hm-pill span { color:var(--gold2) }

/* ─── HOMEPAGE ───────────────────────────────────────── */
#homepage { position:relative; z-index:1 }
.home-intro {
  max-width:1400px; margin:0 auto;
  padding:64px 72px 0;
  display:flex; align-items:baseline; justify-content:space-between; gap:24px;
  flex-wrap:wrap;
}
.home-intro-title {
  font-family:var(--serif); font-size:1.8rem; font-weight:300;
  color:var(--plat2); letter-spacing:.02em;
}
.home-intro-title strong { color:var(--plat); font-weight:400 }
.home-count {
  font-family:var(--mono); font-size:.65rem; letter-spacing:.2em;
  color:var(--gold); text-transform:uppercase;
}
.home-grid {
  max-width:1400px; margin:0 auto;
  padding:40px 72px 100px;
  display:grid;
  grid-template-columns:repeat(auto-fill, minmax(380px, 1fr));
  gap:24px;
}

/* ─── PROJECT CARD ───────────────────────────────────── */
.proj-card {
  background:var(--ink2);
  border:1px solid var(--border);
  border-radius:var(--r2);
  overflow:hidden;
  cursor:pointer;
  position:relative;
  opacity:0;
  transform:translateY(32px);
  transition:border-color var(--t2) var(--ease), box-shadow var(--t2) var(--ease);
}
.proj-card.visible { opacity:1; transform:translateY(0) }
.proj-card:hover { border-color:var(--border2); box-shadow:0 32px 80px rgba(0,0,0,.7), 0 0 0 1px var(--border2), inset 0 1px 0 rgba(255,255,255,.04) }
.proj-card:focus { outline:2px solid var(--gold); outline-offset:3px }

.proj-cover-wrap {
  position:relative; overflow:hidden;
  aspect-ratio:16/9;
  background:var(--ink3);
}
.proj-cover,
.proj-cover-video {
  width:100%; height:100%; object-fit:cover; display:block;
  transition:transform .8s var(--ease);
}
.proj-card:hover .proj-cover,
.proj-card:hover .proj-cover-video { transform:scale(1.08) }
.proj-cover-placeholder {
  width:100%; height:100%;
  display:flex; align-items:center; justify-content:center;
  font-size:2.5rem;
  background:linear-gradient(135deg, var(--ink3), var(--ink4));
}
.proj-cover-gradient {
  position:absolute; inset:0;
  background:linear-gradient(to top, rgba(2,2,9,.92) 0%, rgba(2,2,9,.3) 50%, transparent 100%);
}
.proj-cover-number {
  position:absolute; top:20px; left:20px;
  font-family:var(--mono); font-size:.6rem; letter-spacing:.25em;
  color:rgba(200,165,90,.6);
  background:rgba(2,2,9,.5);
  backdrop-filter:blur(8px);
  border:1px solid var(--border);
  border-radius:999px;
  padding:4px 12px;
}
.proj-cover-cta {
  position:absolute; bottom:20px; right:20px;
  display:flex; align-items:center; gap:8px;
  font-family:var(--mono); font-size:.6rem; letter-spacing:.18em; text-transform:uppercase;
  color:var(--gold2);
  background:rgba(2,2,9,.6); backdrop-filter:blur(12px);
  border:1px solid var(--border2);
  border-radius:999px;
  padding:6px 16px;
  opacity:0; transform:translateY(8px);
  transition:opacity var(--t), transform var(--t);
}
.proj-card:hover .proj-cover-cta { opacity:1; transform:translateY(0) }

.proj-info { padding:24px 28px 28px }
.proj-eyebrow {
  font-family:var(--mono); font-size:.6rem; letter-spacing:.3em; text-transform:uppercase;
  color:var(--gold); margin-bottom:10px;
}
.proj-name {
  font-family:var(--serif);
  font-size:1.5rem; font-weight:400; line-height:1.2;
  color:var(--plat); margin-bottom:16px;
  word-break:break-word;
}
.proj-badges { display:flex; flex-wrap:wrap; gap:6px; margin-bottom:14px }
.badge {
  font-family:var(--mono); font-size:.6rem; letter-spacing:.08em;
  color:var(--plat3); padding:4px 12px;
  border:1px solid var(--border);
  border-radius:999px;
  background:rgba(255,255,255,.02);
  transition:border-color var(--t), color var(--t);
}
.badge-video { border-color:rgba(200,165,90,.3); color:var(--gold2) }
.badge-img   { border-color:var(--border) }
.badge-pdf   { border-color:var(--border) }
.proj-divider { height:1px; background:var(--border); margin-bottom:14px }
.proj-total {
  font-family:var(--mono); font-size:.6rem; letter-spacing:.12em;
  color:var(--plat3);
  display:flex; align-items:center; gap:6px;
}
.proj-total::before {
  content:""; display:inline-block;
  width:6px; height:1px; background:var(--gold);
}

/* ─── PROJECT PAGE ───────────────────────────────────── */
.project-page { position:relative; z-index:1; min-height:100vh }

/* Hero */
.proj-hero {
  position:relative;
  height:100vh; min-height:600px;
  overflow:hidden;
  background:var(--ink);
}
.hero-slides { position:absolute; inset:0 }
.hero-slide {
  position:absolute; inset:0;
  opacity:0;
  transition:opacity 1.2s var(--ease);
}
.hero-slide.active { opacity:1 }
.hero-slide img {
  width:100%; height:100%; object-fit:cover;
  animation:kenBurns 12s var(--ease) forwards;
}
.hero-slide video {
  width:100%; height:100%; object-fit:cover;
}
.hero-gradient {
  position:absolute; inset:0;
  background:linear-gradient(to top, rgba(2,2,9,.95) 0%, rgba(2,2,9,.4) 50%, rgba(2,2,9,.2) 100%),
             linear-gradient(to right, rgba(2,2,9,.3) 0%, transparent 60%);
  z-index:1;
}
.hero-content {
  position:absolute; bottom:0; left:0; right:0;
  z-index:2; padding:0 72px 72px;
  display:flex; align-items:flex-end; justify-content:space-between; gap:32px;
  flex-wrap:wrap;
}
.hero-text {}
.back-btn {
  display:inline-flex; align-items:center; gap:8px;
  font-family:var(--mono); font-size:.65rem; letter-spacing:.15em; text-transform:uppercase;
  color:var(--plat2); background:rgba(2,2,9,.5); backdrop-filter:blur(12px);
  border:1px solid var(--border); border-radius:999px;
  padding:8px 20px; cursor:pointer;
  transition:all var(--t);
  margin-bottom:24px; display:inline-block;
}
.back-btn:hover { color:var(--gold); border-color:var(--border2); background:rgba(200,165,90,.06) }
.hero-eyebrow {
  font-family:var(--mono); font-size:.62rem; letter-spacing:.3em; text-transform:uppercase;
  color:var(--gold); margin-bottom:16px;
}
.hero-title {
  font-family:var(--serif);
  font-size:clamp(2.5rem,6vw,5rem);
  font-weight:300; letter-spacing:-.02em; line-height:1.05;
  color:var(--plat); margin-bottom:24px;
}
.hero-badges { display:flex; flex-wrap:wrap; gap:8px }
.hero-slide-nav {
  display:flex; gap:10px; align-items:flex-end; padding-bottom:4px;
}
.hero-dot {
  width:6px; height:6px; border-radius:50%;
  background:rgba(200,165,90,.3); cursor:pointer;
  transition:background var(--t), width var(--t);
}
.hero-dot.active { width:24px; border-radius:999px; background:var(--gold) }
.hero-counter {
  font-family:var(--mono); font-size:.6rem; letter-spacing:.15em;
  color:var(--plat3); white-space:nowrap;
}
.hero-scroll-hint {
  position:absolute; bottom:72px; left:50%; transform:translateX(-50%);
  z-index:2;
  display:flex; flex-direction:column; align-items:center; gap:8px;
  animation:scrollBounce 2s ease-in-out infinite;
}
.hero-scroll-hint span {
  font-family:var(--mono); font-size:.55rem; letter-spacing:.2em; text-transform:uppercase;
  color:var(--plat3);
}
.hero-scroll-arrow {
  width:20px; height:20px;
  border-right:1px solid var(--plat3);
  border-bottom:1px solid var(--plat3);
  transform:rotate(45deg);
}

/* ─── TAB BAR ────────────────────────────────────────── */
.tab-bar {
  position:sticky; top:0; z-index:50;
  padding:0 24px 0 16px;
  background:rgba(2,2,9,.95);
  backdrop-filter:blur(28px) saturate(1.6);
  border-bottom:1px solid var(--border);
  display:flex; align-items:center; gap:0; overflow-x:auto;
  -webkit-overflow-scrolling:touch;
  box-shadow:0 4px 24px rgba(0,0,0,.4);
}
.tab-bar::-webkit-scrollbar { height:0 }
.tab-bar-back {
  flex-shrink:0;
  display:inline-flex; align-items:center; gap:7px;
  font-family:var(--mono); font-size:.62rem; letter-spacing:.14em; text-transform:uppercase;
  color:var(--gold); background:rgba(200,165,90,.08);
  border:1px solid var(--border2); border-radius:999px;
  padding:7px 16px; cursor:pointer;
  transition:all var(--t); white-space:nowrap;
  margin-right:16px;
}
.tab-bar-back:hover { background:rgba(200,165,90,.18); border-color:var(--border3); color:var(--gold2) }
.tab-bar-back svg { transition:transform var(--t) }
.tab-bar-back:hover svg { transform:translateX(-3px) }
.tab-bar-divider {
  flex-shrink:0; width:1px; height:24px;
  background:var(--border); margin-right:4px;
}
.tab-btn {
  font-family:var(--mono); font-size:.68rem; letter-spacing:.12em; text-transform:uppercase;
  color:var(--plat3); background:transparent; border:none;
  padding:20px 24px; cursor:pointer;
  border-bottom:2px solid transparent;
  white-space:nowrap;
  transition:all var(--t);
}
.tab-btn:hover { color:var(--plat2) }
.tab-btn.active { color:var(--gold2); border-bottom-color:var(--gold) }
.tab-count {
  font-size:.55rem; color:var(--plat3); margin-left:6px;
  background:var(--ink4); padding:2px 7px; border-radius:999px;
}
.tab-btn.active .tab-count { background:rgba(200,165,90,.15); color:var(--gold) }
.tab-content { padding:56px 72px 100px; max-width:1600px; margin:0 auto }

/* ─── GALLERY (MASONRY) ──────────────────────────────── */
.gallery-header {
  display:flex; align-items:baseline; justify-content:space-between;
  margin-bottom:32px; flex-wrap:wrap; gap:16px;
}
.gallery-title {
  font-family:var(--serif); font-size:2rem; font-weight:300; color:var(--plat2);
}
.gallery-title strong { color:var(--plat) }
.gallery-filters { display:flex; gap:8px; flex-wrap:wrap }
.gal-filter {
  font-family:var(--mono); font-size:.6rem; letter-spacing:.12em; text-transform:uppercase;
  color:var(--plat3); background:transparent;
  border:1px solid var(--border); border-radius:999px;
  padding:5px 14px; cursor:pointer;
  transition:all var(--t);
}
.gal-filter:hover { color:var(--plat2); border-color:var(--border2) }
.gal-filter.active { color:var(--gold); border-color:var(--gold); background:var(--gold-dim) }

.gallery {
  columns:3; column-gap:16px;
}
@media(max-width:1100px) { .gallery { columns:2 } }
@media(max-width:600px)  { .gallery { columns:1 } }

.gal-item {
  break-inside:avoid;
  position:relative;
  overflow:hidden;
  border-radius:var(--r);
  background:var(--ink3);
  border:1px solid var(--border);
  cursor:zoom-in;
  margin-bottom:16px;
  opacity:0;
  transform:translateY(20px);
  transition:opacity .6s var(--ease), transform .6s var(--ease),
             border-color var(--t), box-shadow var(--t);
}
.gal-item.visible { opacity:1; transform:translateY(0) }
.gal-item:hover {
  border-color:var(--border2);
  box-shadow:0 20px 60px rgba(0,0,0,.6), 0 0 0 1px var(--border2);
  z-index:2;
}
.gal-item img {
  width:100%; display:block;
  transition:transform .6s var(--ease);
}
.gal-item:hover img { transform:scale(1.04) }
.gal-caption {
  position:absolute; bottom:0; left:0; right:0;
  background:linear-gradient(transparent, rgba(2,2,9,.88));
  padding:32px 14px 12px;
  opacity:0; transition:opacity var(--t);
}
.gal-item:hover .gal-caption { opacity:1 }
.gal-cap-name {
  font-family:var(--mono); font-size:.58rem; letter-spacing:.08em;
  color:rgba(237,232,224,.8); word-break:break-word;
}
.gal-cap-sub {
  font-family:var(--mono); font-size:.55rem; letter-spacing:.06em;
  color:var(--gold); margin-top:2px;
}
.gal-cap-zoom {
  position:absolute; top:12px; right:12px;
  width:32px; height:32px;
  background:rgba(2,2,9,.6); backdrop-filter:blur(8px);
  border:1px solid var(--border2);
  border-radius:50%;
  display:flex; align-items:center; justify-content:center;
  opacity:0; transition:opacity var(--t);
  color:var(--gold2);
}
.gal-item:hover .gal-cap-zoom { opacity:1 }

/* ─── PDF CARDS ──────────────────────────────────────── */
.pdf-grid {
  display:grid;
  grid-template-columns:repeat(auto-fill, minmax(280px, 1fr));
  gap:24px;
}
.pdf-card {
  background:var(--ink2);
  border:1px solid var(--border);
  border-radius:var(--r2);
  overflow:hidden;
  cursor:pointer;
  opacity:0; transform:translateY(20px);
  transition:opacity .6s var(--ease), transform .6s var(--ease),
             border-color var(--t), box-shadow var(--t);
}
.pdf-card.visible { opacity:1; transform:translateY(0) }
.pdf-card:hover {
  border-color:var(--border2);
  box-shadow:0 24px 60px rgba(0,0,0,.65), 0 0 0 1px var(--border2);
}
.pdf-card-preview {
  position:relative;
  aspect-ratio:3/4;
  background:linear-gradient(135deg, var(--ink3), var(--ink4));
  overflow:hidden;
}
.pdf-thumb-canvas {
  width:100%; height:100%;
  object-fit:contain; display:block;
}
.pdf-card-overlay {
  position:absolute; inset:0;
  background:rgba(2,2,9,.4);
  display:flex; flex-direction:column; align-items:center; justify-content:center;
  gap:12px;
  opacity:0; transition:opacity var(--t);
}
.pdf-card:hover .pdf-card-overlay { opacity:1 }
.pdf-open-icon {
  width:56px; height:56px;
  background:rgba(200,165,90,.12);
  backdrop-filter:blur(8px);
  border:1px solid var(--border2);
  border-radius:50%;
  display:flex; align-items:center; justify-content:center;
  color:var(--gold2);
}
.pdf-open-label {
  font-family:var(--mono); font-size:.62rem; letter-spacing:.2em; text-transform:uppercase;
  color:var(--gold2);
}
.pdf-card-info { padding:18px 20px 22px }
.pdf-card-tag {
  font-family:var(--mono); font-size:.58rem; letter-spacing:.2em; text-transform:uppercase;
  color:var(--gold); margin-bottom:8px;
}
.pdf-card-name {
  font-family:var(--serif); font-size:1.1rem; font-weight:400; color:var(--plat);
  line-height:1.3; margin-bottom:8px; word-break:break-word;
}
.pdf-card-meta {
  font-family:var(--mono); font-size:.6rem; letter-spacing:.1em;
  color:var(--plat3);
}

/* ─── PDF MODAL ──────────────────────────────────────── */
#pdf-modal {
  display:none;
  position:fixed; inset:0; z-index:2000;
  background:rgba(2,2,9,.96);
  backdrop-filter:blur(20px);
  flex-direction:column;
}
#pdf-modal.open { display:flex }

.pdf-modal-header {
  display:flex; align-items:center; gap:16px; flex-wrap:wrap;
  padding:16px 28px;
  background:rgba(2,2,9,.9);
  border-bottom:1px solid var(--border);
  position:relative; z-index:1;
}
.pdf-modal-title {
  flex:1; min-width:0;
  font-family:var(--serif); font-size:1.1rem; font-weight:400; color:var(--plat);
  overflow:hidden; text-overflow:ellipsis; white-space:nowrap;
}
.pdf-modal-tag {
  font-family:var(--mono); font-size:.6rem; letter-spacing:.2em; text-transform:uppercase;
  color:var(--gold);
}
.pdf-modal-close {
  width:36px; height:36px; border-radius:50%;
  background:var(--ink3); border:1px solid var(--border);
  color:var(--plat2); font-size:1.1rem;
  cursor:pointer; display:flex; align-items:center; justify-content:center;
  transition:all var(--t); flex-shrink:0; font-family:var(--mono);
}
.pdf-modal-close:hover { border-color:var(--border2); color:var(--gold); background:var(--ink4) }

.pdf-modal-body {
  flex:1; display:flex; overflow:hidden; position:relative;
}
.pdf-main-area {
  flex:1; display:flex; align-items:center; justify-content:center;
  position:relative; overflow:hidden; padding:24px;
}
#pdf-modal-canvas {
  max-width:100%; max-height:100%;
  box-shadow:0 8px 60px rgba(0,0,0,.8);
  border-radius:4px; display:block;
  transition:opacity .25s var(--ease);
}
.pdf-page-loading {
  position:absolute; inset:0;
  display:flex; align-items:center; justify-content:center;
  font-family:var(--mono); font-size:.75rem; letter-spacing:.2em;
  color:var(--plat3);
  background:var(--ink);
  transition:opacity .3s;
}
.pdf-page-loading.hidden { opacity:0; pointer-events:none }

.pdf-nav-btn {
  position:absolute; top:50%; transform:translateY(-50%);
  width:52px; height:52px; border-radius:50%;
  background:rgba(2,2,9,.7); backdrop-filter:blur(12px);
  border:1px solid var(--border2);
  color:var(--gold); font-size:1.4rem; font-family:var(--mono);
  cursor:pointer; display:flex; align-items:center; justify-content:center;
  transition:all var(--t); z-index:5;
}
.pdf-nav-btn:hover { background:rgba(200,165,90,.12); border-color:var(--border3) }
.pdf-nav-btn:disabled { opacity:.25; cursor:default }
#pdf-prev { left:20px }
#pdf-next { right:20px }

.pdf-thumb-strip {
  width:140px; flex-shrink:0;
  background:var(--ink2);
  border-left:1px solid var(--border);
  overflow-y:auto; padding:12px;
  display:flex; flex-direction:column; gap:8px;
}
.pdf-thumb-strip::-webkit-scrollbar { width:3px }
.pdf-thumb-strip::-webkit-scrollbar-thumb { background:var(--border2); border-radius:999px }
.pdf-thumb-item {
  position:relative; cursor:pointer; border-radius:6px; overflow:hidden;
  border:1.5px solid transparent;
  transition:border-color var(--t);
  background:var(--ink3);
}
.pdf-thumb-item.active { border-color:var(--gold) }
.pdf-thumb-item canvas { width:100%; display:block; border-radius:4px }
.pdf-thumb-num {
  position:absolute; bottom:4px; right:6px;
  font-family:var(--mono); font-size:.5rem; letter-spacing:.06em;
  color:rgba(237,232,224,.6);
}

.pdf-modal-footer {
  display:flex; align-items:center; justify-content:center; gap:16px;
  padding:14px 28px;
  background:rgba(2,2,9,.9);
  border-top:1px solid var(--border);
}
.pdf-page-info {
  font-family:var(--mono); font-size:.7rem; letter-spacing:.1em;
  color:var(--gold2); min-width:80px; text-align:center;
}
.pdf-footer-btn {
  font-family:var(--mono); font-size:.65rem; letter-spacing:.1em;
  color:var(--gold); border:1px solid var(--border2);
  background:var(--gold-dim); border-radius:999px;
  padding:6px 18px; cursor:pointer; text-decoration:none;
  transition:all var(--t); white-space:nowrap;
}
.pdf-footer-btn:hover { background:rgba(200,165,90,.18); border-color:var(--border3) }

/* ─── LIGHTBOX ───────────────────────────────────────── */
#lightbox {
  display:none; position:fixed; inset:0; z-index:3000;
  background:rgba(2,2,9,.97);
  backdrop-filter:blur(16px);
  align-items:center; justify-content:center; flex-direction:column;
}
#lightbox.open { display:flex }
#lightbox-wrap {
  position:relative; max-width:92vw; max-height:86vh;
  display:flex; align-items:center; justify-content:center;
}
#lightbox-img {
  max-width:92vw; max-height:86vh;
  object-fit:contain;
  border-radius:6px;
  box-shadow:0 0 100px rgba(0,0,0,.9);
  transition:opacity .3s var(--ease);
  user-select:none;
}
#lightbox-caption {
  position:absolute; bottom:-44px; left:0; right:0;
  font-family:var(--mono); font-size:.65rem; letter-spacing:.1em;
  color:var(--plat3); text-align:center;
}
.lb-close {
  position:absolute; top:24px; right:28px;
  width:40px; height:40px; border-radius:50%;
  background:rgba(2,2,9,.7); backdrop-filter:blur(8px);
  border:1px solid var(--border);
  color:var(--plat2); font-size:1rem; font-family:var(--mono);
  cursor:pointer; display:flex; align-items:center; justify-content:center;
  transition:all var(--t); z-index:10;
}
.lb-close:hover { color:var(--gold); border-color:var(--border2) }
.lb-nav {
  position:absolute; top:50%; transform:translateY(-50%);
  width:52px; height:52px; border-radius:50%;
  background:rgba(2,2,9,.65); backdrop-filter:blur(12px);
  border:1px solid var(--border2);
  color:var(--plat2); font-size:1.5rem; font-family:var(--mono);
  cursor:pointer; display:flex; align-items:center; justify-content:center;
  transition:all var(--t); z-index:10;
}
.lb-nav:hover { color:var(--gold2); border-color:var(--border3); background:rgba(200,165,90,.08) }
#lb-prev { left:-70px }
#lb-next { right:-70px }
#lb-counter {
  position:absolute; top:24px; left:50%; transform:translateX(-50%);
  font-family:var(--mono); font-size:.62rem; letter-spacing:.15em;
  color:var(--plat3);
  background:rgba(2,2,9,.6); backdrop-filter:blur(8px);
  border:1px solid var(--border); border-radius:999px;
  padding:5px 16px;
}

/* ─── FOOTER ─────────────────────────────────────────── */
footer {
  position:relative; z-index:1;
  border-top:1px solid var(--border);
  padding:40px 72px;
  display:flex; justify-content:space-between; align-items:center; flex-wrap:wrap; gap:16px;
  background:linear-gradient(0deg, var(--ink2), transparent);
}
.footer-logo {
  font-family:var(--serif); font-size:1.2rem; font-weight:300;
  color:var(--gold); letter-spacing:.08em;
}
.footer-meta {
  font-family:var(--mono); font-size:.62rem; letter-spacing:.12em;
  color:var(--plat3);
}

/* ─── CARDS (Word/Table etc) ─────────────────────────── */
.card {
  background:var(--ink2); border:1px solid var(--border); border-radius:var(--r);
  padding:28px; margin-bottom:20px;
  transition:border-color var(--t), box-shadow var(--t);
}
.card:hover { border-color:var(--border2); box-shadow:0 12px 40px rgba(0,0,0,.4) }
.card-label {
  font-family:var(--mono); font-size:.6rem; letter-spacing:.25em; text-transform:uppercase;
  color:var(--gold); margin-bottom:12px;
}
.card-title {
  font-family:var(--serif); font-size:1.2rem; font-weight:400; color:var(--plat);
  margin-bottom:16px;
}
.doc-body { color:var(--plat2); font-size:.9rem; line-height:1.7 }
.doc-h1 { font-family:var(--serif); font-size:1.2rem; color:var(--plat); margin:12px 0 6px; font-weight:400 }
.doc-h2 { font-family:var(--serif); font-size:1rem; color:var(--plat2); margin:10px 0 4px }
.doc-p  { margin-bottom:8px }
.table-wrap { overflow-x:auto; margin-top:16px }
table { width:100%; border-collapse:collapse; font-size:.82rem; font-family:var(--mono) }
th {
  padding:10px 14px; text-align:left;
  background:var(--ink3);
  border-bottom:1px solid var(--border2);
  font-size:.6rem; letter-spacing:.1em; text-transform:uppercase; color:var(--gold);
  white-space:nowrap;
}
td { padding:9px 14px; border-bottom:1px solid var(--border); color:var(--plat2); }
tr:last-child td { border-bottom:none }
tr:hover td { background:var(--ink3); color:var(--plat) }
.table-note { font-family:var(--mono); font-size:.62rem; color:var(--plat3); margin-top:10px; letter-spacing:.06em }
.json-pre {
  font-family:var(--mono); font-size:.72rem; color:var(--plat2);
  background:var(--ink3); border:1px solid var(--border);
  border-radius:8px; padding:16px; overflow-x:auto; line-height:1.6;
  max-height:400px; overflow-y:auto;
}
.muted { color:var(--plat3); font-style:italic; font-size:.88rem }
.meta  { font-family:var(--mono); font-size:.68rem; color:var(--plat3); margin-bottom:4px }

/* ─── ANIMATIONS ─────────────────────────────────────── */
@keyframes loaderFadeIn  { to { opacity:1 } }
@keyframes loaderExpand  { to { transform:scaleX(1) } }
@keyframes slideUp       { from { opacity:0; transform:translateY(16px) } to { opacity:1; transform:translateY(0) } }
@keyframes kenBurns      { from { transform:scale(1) } to { transform:scale(1.1) } }
@keyframes orbFloat      { 0%,100%{transform:translate(0,0) scale(1)} 33%{transform:translate(30px,-20px) scale(1.05)} 66%{transform:translate(-20px,15px) scale(.95)} }
@keyframes scrollBounce  { 0%,100%{transform:translateX(-50%) translateY(0)} 50%{transform:translateX(-50%) translateY(8px)} }
@keyframes fadeIn        { from{opacity:0} to{opacity:1} }
@keyframes pageFadeUp    { from{opacity:0;transform:translateY(20px)} to{opacity:1;transform:translateY(0)} }

/* ─── RESPONSIVE ─────────────────────────────────────── */
@media(max-width:1024px) {
  .site-header,.home-intro { padding-left:40px; padding-right:40px }
  .home-grid { padding:40px 40px 80px; grid-template-columns:repeat(auto-fill,minmax(300px,1fr)) }
  .hero-content { padding:0 40px 56px }
  .tab-bar,.tab-content { padding-left:40px; padding-right:40px }
  footer { padding:32px 40px }
}
@media(max-width:640px) {
  .site-header,.home-intro { padding-left:20px; padding-right:20px }
  .home-grid { padding:28px 20px 60px; grid-template-columns:1fr }
  .hero-content { padding:0 20px 40px }
  .hero-title { font-size:2.2rem }
  .tab-bar,.tab-content { padding-left:20px; padding-right:20px }
  .gallery { columns:1 }
  .pdf-grid { grid-template-columns:1fr }
  .pdf-thumb-strip { display:none }
  #lb-prev { left:-10px } #lb-next { right:-10px }
  footer { padding:24px 20px; flex-direction:column; text-align:center }
}
"""


def _html_foot(now: str) -> str:
    return f"""
<!-- ── AMBIENT ORBS ───────────────────────────────── -->
<div class="ambient-orb orb1"></div>
<div class="ambient-orb orb2"></div>
<div class="ambient-orb orb3"></div>

<!-- ── CURSOR ────────────────────────────────────── -->
<div id="cursor"></div>
<div id="cursor-ring"></div>

<!-- ── SCROLL PROGRESS ───────────────────────────── -->
<div id="progress-bar"></div>

<!-- ── PDF MODAL ─────────────────────────────────── -->
<div id="pdf-modal" role="dialog" aria-modal="true">
  <div class="pdf-modal-header">
    <span class="pdf-modal-tag">📄 PDF</span>
    <div class="pdf-modal-title" id="pdf-modal-title">Yükleniyor…</div>
    <button class="pdf-modal-close" onclick="closePdfModal()" aria-label="Kapat">✕</button>
  </div>
  <div class="pdf-modal-body">
    <div class="pdf-main-area">
      <button class="pdf-nav-btn" id="pdf-prev" onclick="pdfModalPrev()" disabled>‹</button>
      <canvas id="pdf-modal-canvas"></canvas>
      <div class="pdf-page-loading" id="pdf-page-loading">Sayfa yükleniyor…</div>
      <button class="pdf-nav-btn" id="pdf-next" onclick="pdfModalNext()">›</button>
    </div>
    <div class="pdf-thumb-strip" id="pdf-thumb-strip"></div>
  </div>
  <div class="pdf-modal-footer">
    <button class="pdf-footer-btn" onclick="pdfModalPrev()">‹ Önceki</button>
    <div class="pdf-page-info" id="pdf-page-info">— / —</div>
    <button class="pdf-footer-btn" onclick="pdfModalNext()">Sonraki ›</button>
    <button class="pdf-footer-btn" onclick="pdfModalFullscreen()">⤢ Tam Ekran</button>
    <a class="pdf-footer-btn" id="pdf-download-btn" href="#" download>⬇ İndir</a>
  </div>
</div>

<!-- ── LIGHTBOX ──────────────────────────────────── -->
<div id="lightbox" role="dialog" aria-modal="true">
  <button class="lb-close" onclick="closeLightbox()" aria-label="Kapat">✕</button>
  <div id="lb-counter">1 / 1</div>
  <div id="lightbox-wrap">
    <button class="lb-nav" id="lb-prev" onclick="lightboxStep(-1)" aria-label="Önceki">‹</button>
    <img id="lightbox-img" src="" alt=""/>
    <button class="lb-nav" id="lb-next" onclick="lightboxStep(1)" aria-label="Sonraki">›</button>
    <div id="lightbox-caption"></div>
  </div>
</div>

<!-- ── FOOTER ────────────────────────────────────── -->
<footer>
  <span class="footer-logo">◈ Proje Sunumu</span>
  <span class="footer-meta">drive_to_html.py · {now}</span>
</footer>

<script>
// ═══════════════════════════════════════════════════
//  LOADER
// ═══════════════════════════════════════════════════
window.addEventListener('load', function() {{
  setTimeout(function() {{
    document.getElementById('loader').classList.add('hidden');
    document.getElementById('loader').addEventListener('transitionend', function() {{
      this.remove();
    }}, {{once:true}});
  }}, 800);
}});

// ═══════════════════════════════════════════════════
//  CUSTOM CURSOR
// ═══════════════════════════════════════════════════
(function() {{
  var cur = document.getElementById('cursor');
  var ring = document.getElementById('cursor-ring');
  var mx=0,my=0,rx=0,ry=0;
  document.addEventListener('mousemove', function(e) {{
    mx=e.clientX; my=e.clientY;
    cur.style.left=mx+'px'; cur.style.top=my+'px';
  }});
  function animRing() {{
    rx += (mx-rx)*0.12; ry += (my-ry)*0.12;
    ring.style.left=rx+'px'; ring.style.top=ry+'px';
    requestAnimationFrame(animRing);
  }}
  animRing();
}})();

// ═══════════════════════════════════════════════════
//  SCROLL PROGRESS
// ═══════════════════════════════════════════════════
window.addEventListener('scroll', function() {{
  var s=document.documentElement;
  var pct=(s.scrollTop/(s.scrollHeight-s.clientHeight))*100;
  document.getElementById('progress-bar').style.width=pct+'%';
}}, {{passive:true}});

// ═══════════════════════════════════════════════════
//  INTERSECTION OBSERVER — reveal on scroll
// ═══════════════════════════════════════════════════
var _observer = new IntersectionObserver(function(entries) {{
  entries.forEach(function(e) {{
    if(e.isIntersecting) {{
      e.target.classList.add('visible');
      _observer.unobserve(e.target);
    }}
  }});
}}, {{threshold:0.05, rootMargin:'0px 0px -40px 0px'}});

function observeItems() {{
  document.querySelectorAll('.proj-card:not(.visible),.gal-item:not(.visible),.pdf-card:not(.visible)').forEach(function(el,i) {{
    el.style.transitionDelay=(i*0.05)+'s';
    _observer.observe(el);
  }});
}}
document.addEventListener('DOMContentLoaded', observeItems);

// ── Lazy video: play on card hover ─────────────────
(function() {{
  document.addEventListener('mouseover', function(e) {{
    var card=e.target.closest('.proj-card');
    if(!card) return;
    var v=card.querySelector('[data-lazy-video]');
    if(v && v.paused) {{ v.play().catch(function(){{}}); }}
  }});
  document.addEventListener('mouseout', function(e) {{
    var card=e.target.closest('.proj-card');
    if(!card) return;
    if(e.relatedTarget && card.contains(e.relatedTarget)) return;
    var v=card.querySelector('[data-lazy-video]');
    if(v && !v.paused) {{ v.pause(); v.currentTime=0; }}
  }});
}})();

// ═══════════════════════════════════════════════════
//  PROJECT NAVIGATION
// ═══════════════════════════════════════════════════
function openProjectAnim(cardEl, slug) {{
  var rect=cardEl.getBoundingClientRect();
  var ov=document.createElement('div');
  var cx=rect.left+rect.width/2, cy=rect.top+rect.height/2;
  ov.style.cssText='position:fixed;left:'+cx+'px;top:'+cy+'px;width:6px;height:6px;border-radius:50%;background:var(--ink2);transform:translate(-50%,-50%) scale(0);transition:transform .6s cubic-bezier(.4,0,.2,1),opacity .2s .5s;z-index:900;pointer-events:none;';
  document.body.appendChild(ov);
  var maxR=Math.sqrt(Math.pow(window.innerWidth,2)+Math.pow(window.innerHeight,2));
  requestAnimationFrame(function() {{
    ov.style.transform='translate(-50%,-50%) scale('+maxR+')';
  }});
  setTimeout(function() {{
    openProject(slug);
    ov.style.opacity='0';
    setTimeout(function(){{ov.remove();}},220);
  }}, 520);
}}

function openProject(slug) {{
  // pause all homepage card videos
  document.querySelectorAll('.proj-cover-video').forEach(function(v){{v.pause();}});
  document.getElementById('homepage').style.display='none';
  document.querySelectorAll('.project-page').forEach(function(p){{p.style.display='none';}});
  var pg=document.getElementById('page-'+slug);
  if(pg) {{
    pg.style.display='block';
    pg.style.animation='pageFadeUp .4s var(--ease) both';
    window.scrollTo(0,0);
    startHeroSlideshow(slug);
    setTimeout(observeItems, 100);
    initPdfThumbs(pg);
  }}
}}

function closeProject() {{
  stopHeroSlideshow();
  document.querySelectorAll('.project-page').forEach(function(p){{p.style.display='none';}});
  // also pause any hero videos
  document.querySelectorAll('.hero-slide video').forEach(function(v){{v.pause();}});
  var hp=document.getElementById('homepage');
  hp.style.display='block';
  hp.style.animation='pageFadeUp .4s var(--ease) both';
  window.scrollTo(0,0);
  setTimeout(observeItems,50);
}}

function switchTab(slug, key, btn) {{
  var page=document.getElementById('page-'+slug);
  if(!page) return;
  page.querySelectorAll('.tab-panel').forEach(function(p){{p.style.display='none';}});
  page.querySelectorAll('.tab-btn').forEach(function(b){{b.classList.remove('active');}});
  var panel=document.getElementById('panel-'+slug+'-'+key);
  if(panel) {{ panel.style.display=''; setTimeout(function(){{observeItems();}},50); }}
  if(btn) btn.classList.add('active');
}}

// ═══════════════════════════════════════════════════
//  HERO SLIDESHOW (Ken Burns)
// ═══════════════════════════════════════════════════
var _heroTimer=null, _heroIdx=0, _heroSlides=[];
function startHeroSlideshow(slug) {{
  var hero=document.querySelector('#page-'+slug+' .hero-slides');
  if(!hero) return;
  _heroSlides=Array.from(hero.querySelectorAll('.hero-slide'));
  if(!_heroSlides.length) return;
  _heroIdx=0;
  _heroSlides.forEach(function(s){{s.classList.remove('active');}});
  _heroSlides[0].classList.add('active');
  updateHeroDots(slug,0);
  if(_heroSlides.length>1) {{
    _heroTimer=setInterval(function() {{
      _heroSlides[_heroIdx].classList.remove('active');
      _heroIdx=(_heroIdx+1)%_heroSlides.length;
      _heroSlides[_heroIdx].classList.add('active');
      // restart Ken Burns
      var img=_heroSlides[_heroIdx].querySelector('img');
      if(img) {{ img.style.animation='none'; void img.offsetWidth; img.style.animation='kenBurns 12s var(--ease) forwards'; }}
      updateHeroDots(slug,_heroIdx);
    }}, 5000);
  }}
}}
function stopHeroSlideshow() {{
  if(_heroTimer) {{ clearInterval(_heroTimer); _heroTimer=null; }}
}}
function goHeroSlide(slug, idx) {{
  if(_heroTimer) clearInterval(_heroTimer);
  _heroSlides.forEach(function(s){{s.classList.remove('active');}});
  _heroIdx=idx;
  _heroSlides[_heroIdx].classList.add('active');
  var img=_heroSlides[_heroIdx].querySelector('img');
  if(img) {{ img.style.animation='none'; void img.offsetWidth; img.style.animation='kenBurns 12s var(--ease) forwards'; }}
  updateHeroDots(slug,idx);
  if(_heroSlides.length>1) {{
    _heroTimer=setInterval(function() {{
      _heroSlides[_heroIdx].classList.remove('active');
      _heroIdx=(_heroIdx+1)%_heroSlides.length;
      _heroSlides[_heroIdx].classList.add('active');
      var img2=_heroSlides[_heroIdx].querySelector('img');
      if(img2) {{ img2.style.animation='none'; void img2.offsetWidth; img2.style.animation='kenBurns 12s var(--ease) forwards'; }}
      updateHeroDots(slug,_heroIdx);
    }}, 5000);
  }}
}}
function updateHeroDots(slug, idx) {{
  var dotsEl=document.querySelector('#page-'+slug+' .hero-slide-nav');
  if(!dotsEl) return;
  dotsEl.querySelectorAll('.hero-dot').forEach(function(d,i){{
    d.classList.toggle('active',i===idx);
  }});
}}

// ═══════════════════════════════════════════════════
//  GALLERY FILTERS
// ═══════════════════════════════════════════════════
function galFilter(btn, slug, cat) {{
  var bar=btn.closest('.gallery-filters');
  bar.querySelectorAll('.gal-filter').forEach(function(b){{b.classList.remove('active');}});
  btn.classList.add('active');
  var gallery=document.getElementById('gal-'+slug);
  if(!gallery) return;
  gallery.querySelectorAll('.gal-item').forEach(function(item) {{
    if(cat==='all' || item.dataset.cat===cat) {{
      item.style.display='';
    }} else {{
      item.style.display='none';
    }}
  }});
}}

// ═══════════════════════════════════════════════════
//  PDF CARD THUMBNAILS (lazy render page 1)
// ═══════════════════════════════════════════════════
function initPdfThumbs(container) {{
  if(!window.pdfjsLib) {{ setTimeout(function(){{initPdfThumbs(container);}},150); return; }}
  window.pdfjsLib.GlobalWorkerOptions.workerSrc='https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.worker.min.js';
  container.querySelectorAll('.pdf-thumb-canvas[data-pdf]').forEach(function(canvas) {{
    var url=canvas.dataset.pdf;
    if(canvas.dataset.loaded) return;
    canvas.dataset.loaded='1';
    window.pdfjsLib.getDocument(url).promise.then(function(pdf) {{
      pdf.getPage(1).then(function(page) {{
        var vp=page.getViewport({{scale:1}});
        var scale=canvas.parentElement.offsetWidth/vp.width;
        var vp2=page.getViewport({{scale:scale*0.9}});
        canvas.width=vp2.width; canvas.height=vp2.height;
        canvas.style.width='100%'; canvas.style.height='auto';
        page.render({{canvasContext:canvas.getContext('2d'),viewport:vp2}});
      }});
    }}).catch(function(){{ canvas.style.display='none'; }});
  }});
}}

// ═══════════════════════════════════════════════════
//  PDF MODAL
// ═══════════════════════════════════════════════════
var _pdfModal = {{
  doc:null, cur:1, total:0, url:'', name:''
}};
function openPdfModal(url, name, totalHint) {{
  if(!window.pdfjsLib) {{ alert('PDF görüntüleyici yükleniyor, lütfen bir dakika bekleyin.'); return; }}
  window.pdfjsLib.GlobalWorkerOptions.workerSrc='https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.worker.min.js';
  _pdfModal.url=url; _pdfModal.name=name; _pdfModal.cur=1;
  document.getElementById('pdf-modal-title').textContent=name;
  document.getElementById('pdf-download-btn').href=url;
  document.getElementById('pdf-download-btn').download=name;
  document.getElementById('pdf-modal').classList.add('open');
  document.body.style.overflow='hidden';
  document.getElementById('pdf-page-loading').classList.remove('hidden');
  document.getElementById('pdf-modal-canvas').style.opacity='0';
  document.getElementById('pdf-thumb-strip').innerHTML='';
  window.pdfjsLib.getDocument(url).promise.then(function(pdf) {{
    _pdfModal.doc=pdf; _pdfModal.total=pdf.numPages;
    renderPdfPage(1);
    buildPdfThumbs();
  }}).catch(function(e) {{
    document.getElementById('pdf-page-loading').textContent='PDF açılamadı: '+e.message;
  }});
}}
function renderPdfPage(num) {{
  if(!_pdfModal.doc) return;
  _pdfModal.cur=num;
  document.getElementById('pdf-page-loading').classList.remove('hidden');
  document.getElementById('pdf-modal-canvas').style.opacity='0';
  _pdfModal.doc.getPage(num).then(function(page) {{
    var area=document.querySelector('.pdf-main-area');
    var maxW=area.clientWidth-120, maxH=area.clientHeight-40;
    var vp0=page.getViewport({{scale:1}});
    var scale=Math.min(maxW/vp0.width, maxH/vp0.height, 2.5);
    var vp=page.getViewport({{scale:scale}});
    var canvas=document.getElementById('pdf-modal-canvas');
    canvas.width=vp.width; canvas.height=vp.height;
    page.render({{canvasContext:canvas.getContext('2d'),viewport:vp}}).promise.then(function() {{
      canvas.style.opacity='1';
      document.getElementById('pdf-page-loading').classList.add('hidden');
      document.getElementById('pdf-page-info').textContent=num+' / '+_pdfModal.total;
      document.getElementById('pdf-prev').disabled=(num<=1);
      document.getElementById('pdf-next').disabled=(num>=_pdfModal.total);
      // highlight thumb
      document.querySelectorAll('.pdf-thumb-item').forEach(function(t,i) {{
        t.classList.toggle('active',i+1===num);
      }});
    }});
  }});
}}
function buildPdfThumbs() {{
  var strip=document.getElementById('pdf-thumb-strip');
  strip.innerHTML='';
  var n=Math.min(_pdfModal.total, 50);
  for(var i=1;i<=n;i++) {{
    (function(pnum) {{
      var div=document.createElement('div');
      div.className='pdf-thumb-item'+(pnum===1?' active':'');
      div.onclick=(function(p){{return function(){{renderPdfPage(p);}}}})(pnum);
      var c=document.createElement('canvas');
      var num=document.createElement('span');
      num.className='pdf-thumb-num'; num.textContent=pnum;
      div.appendChild(c); div.appendChild(num); strip.appendChild(div);
      _pdfModal.doc.getPage(pnum).then(function(page) {{
        var vp=page.getViewport({{scale:1}});
        var scale=110/vp.width;
        var vp2=page.getViewport({{scale:scale}});
        c.width=vp2.width; c.height=vp2.height;
        c.style.width='100%'; c.style.height='auto';
        page.render({{canvasContext:c.getContext('2d'),viewport:vp2}});
      }});
    }})(i);
  }}
}}
function pdfModalPrev() {{ if(_pdfModal.cur>1) renderPdfPage(_pdfModal.cur-1); }}
function pdfModalNext() {{ if(_pdfModal.cur<_pdfModal.total) renderPdfPage(_pdfModal.cur+1); }}
function closePdfModal() {{
  document.getElementById('pdf-modal').classList.remove('open');
  document.body.style.overflow='';
  _pdfModal.doc=null;
}}
function pdfModalFullscreen() {{
  var el=document.getElementById('pdf-modal');
  if(el.requestFullscreen) el.requestFullscreen();
  else if(el.webkitRequestFullscreen) el.webkitRequestFullscreen();
}}

// ═══════════════════════════════════════════════════
//  LIGHTBOX (with touch/swipe)
// ═══════════════════════════════════════════════════
var _lbItems=[],_lbIdx=0,_lbTouchX=null;
function openLightbox(el) {{
  var gallery=el.closest('.gallery');
  _lbItems=Array.from(gallery.querySelectorAll('.gal-item:not([style*="none"])'));
  _lbIdx=_lbItems.indexOf(el);
  showLightboxItem();
  document.getElementById('lightbox').classList.add('open');
  document.body.style.overflow='hidden';
}}
function showLightboxItem() {{
  var item=_lbItems[_lbIdx];
  var img=item.querySelector('img');
  var lb=document.getElementById('lightbox-img');
  lb.style.opacity='0';
  lb.onload=function(){{lb.style.opacity='1';}};
  lb.src=img.dataset.large||img.src;
  lb.alt=img.alt;
  var capName=item.querySelector('.gal-cap-name');
  var capSub=item.querySelector('.gal-cap-sub');
  var caption='';
  if(capName) caption+=capName.textContent;
  if(capSub) caption+=' · '+capSub.textContent;
  document.getElementById('lightbox-caption').textContent=caption;
  document.getElementById('lb-counter').textContent=(_lbIdx+1)+' / '+_lbItems.length;
  document.getElementById('lb-prev').style.display=_lbItems.length>1?'':'none';
  document.getElementById('lb-next').style.display=_lbItems.length>1?'':'none';
}}
function lightboxStep(dir) {{
  _lbIdx=(_lbIdx+dir+_lbItems.length)%_lbItems.length;
  showLightboxItem();
}}
function closeLightbox() {{
  document.getElementById('lightbox').classList.remove('open');
  document.body.style.overflow='';
}}
document.getElementById('lightbox').addEventListener('click',function(e){{if(e.target===this)closeLightbox();}});
document.getElementById('lightbox').addEventListener('touchstart',function(e){{_lbTouchX=e.touches[0].clientX;}},{{passive:true}});
document.getElementById('lightbox').addEventListener('touchend',function(e){{
  if(_lbTouchX===null) return;
  var dx=e.changedTouches[0].clientX-_lbTouchX;
  if(Math.abs(dx)>50) lightboxStep(dx<0?1:-1);
  _lbTouchX=null;
}});

// ═══════════════════════════════════════════════════
//  KEYBOARD SHORTCUTS
// ═══════════════════════════════════════════════════
document.addEventListener('keydown', function(e) {{
  var lbOpen=document.getElementById('lightbox').classList.contains('open');
  var pdfOpen=document.getElementById('pdf-modal').classList.contains('open');
  if(e.key==='Escape') {{
    if(lbOpen) closeLightbox();
    else if(pdfOpen) closePdfModal();
    else closeProject();
  }}
  if(lbOpen) {{
    if(e.key==='ArrowRight') lightboxStep(1);
    if(e.key==='ArrowLeft')  lightboxStep(-1);
  }}
  if(pdfOpen) {{
    if(e.key==='ArrowRight'||e.key==='ArrowDown') pdfModalNext();
    if(e.key==='ArrowLeft'||e.key==='ArrowUp')   pdfModalPrev();
  }}
}});
</script>
</body>
</html>
"""


# ════════════════════════════════════════════════════════════
#  10. ANA FONKSİYON
# ════════════════════════════════════════════════════════════
def main():
    print("─" * 52)
    print("  drive_to_html.py — Lüks İmmersive HTML Üretici v2")
    print("─" * 52)

    print("\n[1/4] Google Drive'a bağlanılıyor…")
    Path(ASSETS_DIR).mkdir(exist_ok=True)
    service = get_service()
    print("  ✓ Bağlantı başarılı")

    print(f"\n[2/4] Klasör taranıyor (id: {FOLDER_ID})…")
    files = list_files(service, FOLDER_ID)
    if not files:
        print("  ⚠  Klasörde dosya bulunamadı.")
        return
    print(f"  ✓ {len(files)} dosya bulundu")

    current_hash = compute_manifest(files)
    prev_hash    = load_manifest()
    html_exists  = Path(OUTPUT_FILE).exists()

    if current_hash == prev_hash and html_exists:
        print(f"\n  ✅ Drive'da değişiklik yok — {OUTPUT_FILE} zaten güncel.")
        print(f"     (Yeniden üretmek için '{MANIFEST_FILE}' dosyasını sil.)")
        print("─" * 52)
        return
    elif prev_hash and current_hash != prev_hash:
        print(f"  🔄 Değişiklik tespit edildi — HTML yeniden üretiliyor…")
    else:
        print(f"  🆕 İlk çalıştırma — HTML üretiliyor…")

    print("\n[3/4] Dosyalar işleniyor…")
    stats = cache_stats()
    if stats["files"] > 0:
        print(f"  ℹ  Önbellek: {stats['files']} dosya, {stats['size_mb']:.1f} MB  ({CACHE_DIR}/)")
    else:
        print(f"  ℹ  Önbellek boş — ilk çalıştırmada tüm dosyalar indirilecek")
    _cache_hits = [0]
    _cache_miss = [0]

    from collections import OrderedDict
    projects = OrderedDict()

    def get_project(fp: str) -> str:
        parts = fp.split(" / ")
        return parts[0] if parts else "DİĞER"

    for f in files:
        pname = get_project(f.get("folder_path", ""))
        if pname not in projects:
            projects[pname] = {
                "images":  [],
                "pdfs":    [],
                "tables":  [],
                "words":   [],
                "others":  [],
                "cover":   "",
                "video":   "",
                "total":   0,
                "cover_is_exterior": False,
            }
        projects[pname]["total"] += 1

    total = len(files)
    for i, f in enumerate(files, 1):
        name  = f["name"]
        mime  = f.get("mimeType", "")
        fid   = f["id"]
        tkey  = file_type_key(name, mime)
        fp    = f.get("folder_path", "")
        pname = get_project(fp)
        sub   = " / ".join(fp.split(" / ")[1:]) if " / " in fp else ""
        print(f"  [{i:03}/{total:03}] {name[:48]:<48} [{tkey}]")

        try:
            _prev_files = cache_stats()["files"]
            data, real_mime = download_cached(service, fid, mime, f.get("modifiedTime",""))
            if cache_stats()["files"] > _prev_files: _cache_miss[0] += 1
            else: _cache_hits[0] += 1
            tkey2 = file_type_key(name, real_mime) if real_mime != mime else tkey
            name_up = name.upper()
            asset_uid = hashlib.md5((fid + f.get("modifiedTime","")).encode()).hexdigest()[:12]

            if tkey2 == "other":
                is_video  = (real_mime.startswith("video/") or name.lower().endswith(".mp4"))
                is_tanitim = "TANITIM" in name_up
                if is_video and is_tanitim and not projects[pname]["video"]:
                    vpath = save_video(data, asset_uid)
                    if vpath:
                        projects[pname]["video"] = vpath
                        print(f"    ✓ Video kaydedildi: {name}")
                del data
                continue

            if tkey2 in ("word", "table", "json"):
                del data
                continue
            if tkey2 == "pdf":
                if not any(k in name_up for k in (
                    "SUNUM", "FİYAT", "FIYAT", "ÖDEME", "ODEME",
                    "LİSTE", "LISTE", "TANITIM", "KATALOG", "OTURUM", "PLAN", "BRAVO"
                )):
                    del data
                    continue

            if tkey2 == "image":
                thumb = make_thumb(data, real_mime, asset_uid)
                large = make_large(data, real_mime, asset_uid)
                del data
                is_exterior = "DIŞ CEPHE" in fp.upper() or "DIS CEPHE" in fp.upper()
                projects[pname]["images"].append((thumb, large, html.escape(name), html.escape(sub), html.escape(fp)))
                if not projects[pname]["cover"] and thumb:
                    projects[pname]["cover"] = thumb
                    projects[pname]["cover_is_exterior"] = is_exterior
                elif is_exterior and not projects[pname].get("cover_is_exterior") and thumb:
                    projects[pname]["cover"] = thumb
                    projects[pname]["cover_is_exterior"] = True

            elif tkey2 == "pdf":
                lbl  = "Slides" if mime == "application/vnd.google-apps.presentation" else "PDF"
                card = process_pdf(data, name, lbl)
                del data
                projects[pname]["pdfs"].append(card)

        except Exception as e:
            print(f"    ⚠  İşlenemedi: {e}")
            try: del data
            except: pass

    print(f"\n  Önbellek özeti: {_cache_hits[0]} hit (atlandı), {_cache_miss[0]} miss (indirildi)")
    print(f"[4/4] HTML dosyası yazılıyor → {OUTPUT_FILE}")
    now    = datetime.now().strftime("%d.%m.%Y %H:%M")
    n_proj = len(projects)
    _words = PROJE_ADI.split()
    proje_adi_h1 = " ".join(_words[:-1]) + (f" <em>{_words[-1]}</em>" if len(_words) > 1 else _words[0])

    HTML_HEAD = _html_head(PROJE_ADI)
    HTML_CSS  = _html_css()
    HTML_FOOT = _html_foot(now)

    with open(OUTPUT_FILE, "w", encoding="utf-8") as fh:

        # HEAD + CSS
        fh.write(HTML_HEAD)
        fh.write(HTML_CSS)
        fh.write("</style>\n</head>\n<body>\n\n")

        # LOADER
        fh.write(f"""<div id="loader">
  <div class="loader-logo">{html.escape(PROJE_ADI)}</div>
  <div class="loader-line"></div>
  <div class="loader-sub">Gayrimenkul Portföyü</div>
</div>\n\n""")

        # HEADER
        fh.write(f"""<header class="site-header" id="site-header">
  <div class="header-inner">
    <p class="header-eyebrow">Gayrimenkul Portföyü · Google Drive Arşivi</p>
    <h1>{proje_adi_h1}</h1>
    <p class="header-sub">{html.escape(PROJE_ALT_BASLIK)}</p>
    <div class="header-line"></div>
    <div class="header-meta">
      <div class="hm-pill">Proje <span>{n_proj}</span></div>
      <div class="hm-pill">Toplam Dosya <span>{total}</span></div>
      <div class="hm-pill">Oluşturulma <span>{now}</span></div>
    </div>
  </div>
</header>\n\n""")

        # HOMEPAGE
        fh.write('<div id="homepage">\n')
        fh.write(f'<div class="home-intro">'
                 f'<div class="home-intro-title">Tüm <strong>Projeler</strong></div>'
                 f'<div class="home-count">{n_proj} Proje</div>'
                 f'</div>\n')
        fh.write('<div class="home-grid">\n')

        for idx, (pname, pdata) in enumerate(projects.items(), 1):
            slug      = slugify(pname)
            cover     = pdata["cover"] or ""
            n_img     = len(pdata["images"])
            n_pdf     = len(pdata["pdfs"])
            total_p   = pdata["total"]
            video_src = pdata.get("video", "")
            has_video = bool(video_src)
            num_label = f"{idx:02d}"

            if has_video:
                cover_html = (f'<video class="proj-cover proj-cover-video" '
                              f'src="{video_src}" muted loop playsinline preload="none" '
                              f'poster="{cover}" data-lazy-video="1"></video>')
            elif cover:
                cover_html = f'<img src="{cover}" class="proj-cover" alt="{html.escape(pname)}" loading="lazy"/>'
            else:
                cover_html = '<div class="proj-cover-placeholder">◈</div>'

            badges = ""
            if has_video: badges += '<span class="badge badge-video">▶ Video</span>'
            if n_img:     badges += f'<span class="badge badge-img">🖼 {n_img}</span>'
            if n_pdf:     badges += f'<span class="badge badge-pdf">📄 {n_pdf} PDF</span>'
            cta = "▶ İzle →" if has_video else "Keşfet →"

            fh.write(f"""<div class="proj-card" onclick="openProjectAnim(this,'{slug}')" tabindex="0" role="button"
     onkeydown="if(event.key==='Enter')openProjectAnim(this,'{slug}')"
     aria-label="{html.escape(pname)}">
  <div class="proj-cover-wrap">
    {cover_html}
    <div class="proj-cover-gradient"></div>
    <div class="proj-cover-number">{num_label}</div>
    <div class="proj-cover-cta">{cta}</div>
  </div>
  <div class="proj-info">
    <div class="proj-eyebrow">Proje</div>
    <h2 class="proj-name">{html.escape(pname)}</h2>
    <div class="proj-badges">{badges}</div>
    <div class="proj-divider"></div>
    <div class="proj-total">{total_p} dosya</div>
  </div>
</div>\n""")

        fh.write('</div>\n</div>\n\n<!-- PROJECT PAGES -->\n')

        # PROJECT PAGES
        for pname, pdata in projects.items():
            slug    = slugify(pname)
            n_img   = len(pdata["images"])
            n_pdf   = len(pdata["pdfs"])
            total_p = pdata["total"]
            has_video = bool(pdata.get("video"))

            # Hero images: up to 8 best images for slideshow
            hero_imgs = [item for item in pdata["images"] if item[1]][:8]
            hero_video = pdata.get("video", "")

            badges_h = ""
            if has_video: badges_h += '<span class="badge badge-video">▶ Video</span>'
            if n_img:     badges_h += f'<span class="badge">🖼 {n_img} görsel</span>'
            if n_pdf:     badges_h += f'<span class="badge">📄 {n_pdf} PDF</span>'

            # Build hero slides
            hero_slides_html = ""
            if hero_video:
                hero_slides_html += (f'<div class="hero-slide active">'
                                     f'<video src="{hero_video}" autoplay muted loop playsinline style="width:100%;height:100%;object-fit:cover"></video>'
                                     f'</div>\n')
                for item in hero_imgs:
                    large = item[1]
                    hero_slides_html += f'<div class="hero-slide"><img src="{large}" alt="" loading="lazy"/></div>\n'
            elif hero_imgs:
                for i, item in enumerate(hero_imgs):
                    large  = item[1]
                    active = ' active' if i == 0 else ''
                    hero_slides_html += f'<div class="hero-slide{active}"><img src="{large}" alt="" loading="lazy"/></div>\n'
            else:
                hero_slides_html = '<div class="hero-slide active" style="background:var(--ink3)"></div>\n'

            # Hero dots
            n_slides = (1 if hero_video else 0) + len(hero_imgs)
            dots_html = ""
            for i in range(min(n_slides, 8)):
                active = ' active' if i == 0 else ''
                dots_html += f'<div class="hero-dot{active}" onclick="goHeroSlide(\'{slug}\',{i})"></div>\n'

            # Tab bar (with back button)
            tabs_html = (f'<button class="tab-bar-back" onclick="closeProject()" aria-label="Tüm Projeler">'
                         f'<svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><polyline points="15 18 9 12 15 6"/></svg>'
                         f'Tüm Projeler</button>'
                         f'<div class="tab-bar-divider"></div>')
            if pdata["images"]:
                tabs_html += (f'<button class="tab-btn active" '
                              f'onclick="switchTab(\'{slug}\',\'images\',this)">🖼 Görseller '
                              f'<span class="tab-count">{n_img}</span></button>')
            if pdata["pdfs"]:
                active2 = "" if pdata["images"] else " active"
                tabs_html += (f'<button class="tab-btn{active2}" '
                              f'onclick="switchTab(\'{slug}\',\'pdfs\',this)">📄 Sunum &amp; Fiyat '
                              f'<span class="tab-count">{n_pdf}</span></button>')

            fh.write(f"""<section class="project-page" id="page-{slug}" style="display:none">
  <div class="proj-hero">
    <div class="hero-slides">{hero_slides_html}</div>
    <div class="hero-gradient"></div>
    <div class="hero-content">
      <div class="hero-text">
        <button class="back-btn" onclick="closeProject()">← Tüm Projeler</button>
        <div class="hero-eyebrow">Proje Detayı</div>
        <h1 class="hero-title">{html.escape(pname)}</h1>
        <div class="hero-badges">{badges_h}</div>
      </div>
      <div style="display:flex;flex-direction:column;align-items:flex-end;gap:12px;padding-bottom:4px">
        <div class="hero-slide-nav">{dots_html}</div>
        <div class="hero-counter">{n_slides} görsel</div>
      </div>
    </div>
  </div>
  <div class="tab-bar" id="tabs-{slug}">{tabs_html}</div>
  <div class="tab-content">\n""")

            # IMAGES PANEL
            if pdata["images"]:
                # Collect unique sub-categories for filter
                cats = []
                seen_cats = set()
                for item in pdata["images"]:
                    sub = item[3]  # html-escaped sub-folder
                    if sub and sub not in seen_cats:
                        seen_cats.add(sub)
                        cats.append(sub)

                filters_html = '<button class="gal-filter active" onclick="galFilter(this,\''+slug+'\',\'all\')">Tümü</button>'
                for cat in cats[:6]:
                    short_cat = cat.split(" / ")[-1] if " / " in cat else cat
                    filters_html += f'<button class="gal-filter" onclick="galFilter(this,\'{slug}\',\'{cat}\')">{short_cat}</button>'

                fh.write(f'<div class="tab-panel" id="panel-{slug}-images">\n')
                fh.write(f'<div class="gallery-header">'
                         f'<div class="gallery-title"><strong>{html.escape(pname)}</strong> · Görseller</div>'
                         f'<div class="gallery-filters">{filters_html}</div>'
                         f'</div>\n')
                fh.write(f'<div class="gallery" id="gal-{slug}">\n')
                for item in pdata["images"]:
                    thumb, large, iname, isub, ifp = item[0], item[1], item[2], item[3], item[4]
                    if not thumb:
                        continue
                    cat_attr = f' data-cat="{isub}"' if isub else ''
                    zoom_icon = ('<svg width="14" height="14" viewBox="0 0 24 24" fill="none" '
                                 'stroke="currentColor" stroke-width="2"><circle cx="11" cy="11" r="8"/>'
                                 '<line x1="21" y1="21" x2="16.65" y2="16.65"/>'
                                 '<line x1="11" y1="8" x2="11" y2="14"/>'
                                 '<line x1="8" y1="11" x2="14" y2="11"/></svg>')
                    sub_div = f'<div class="gal-cap-sub">{isub}</div>' if isub else ''
                    fh.write(f'<div class="gal-item" onclick="openLightbox(this)"{cat_attr}>'
                             f'<img src="{thumb}" data-large="{large}" alt="{iname}" loading="lazy"/>'
                             f'<div class="gal-caption">'
                             f'<div class="gal-cap-name">{iname}</div>'
                             f'{sub_div}'
                             f'</div>'
                             f'<div class="gal-cap-zoom">{zoom_icon}</div>'
                             f'</div>\n')
                fh.write('</div>\n</div>\n')  # gallery + tab-panel

            # PDF PANEL
            if pdata["pdfs"]:
                hidden = ' style="display:none"' if pdata["images"] else ""
                fh.write(f'<div class="tab-panel" id="panel-{slug}-pdfs"{hidden}>\n')
                fh.write(f'<div class="gallery-header">'
                         f'<div class="gallery-title"><strong>{html.escape(pname)}</strong> · Belgeler</div>'
                         f'</div>\n')
                fh.write('<div class="pdf-grid">\n')
                for card in pdata["pdfs"]:
                    fh.write(card + "\n")
                fh.write('</div>\n</div>\n')

            fh.write('  </div>\n</section>\n\n')

        # FOOT
        fh.write(HTML_FOOT)

    size_kb = Path(OUTPUT_FILE).stat().st_size / 1024
    asset_files = len(list(Path(ASSETS_DIR).iterdir())) if Path(ASSETS_DIR).exists() else 0
    save_manifest(current_hash)
    print(f"  ✓ {OUTPUT_FILE} oluşturuldu ({size_kb:.0f} KB)")
    print(f"  ✓ {ASSETS_DIR}/ klasörü: {asset_files} dosya")
    print(f"\n  ⚠  HTML ve '{ASSETS_DIR}/' klasörünü birlikte tut!")
    print(f"\n  Tarayıcıda aç:")
    print(f"    python -m http.server 8080  →  http://localhost:8080/{OUTPUT_FILE}")
    print("─" * 52)


if __name__ == "__main__":
    main()